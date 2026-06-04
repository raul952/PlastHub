# -*- coding: utf-8 -*-
import io,sys; sys.stdout=io.TextIOWrapper(sys.stdout.buffer,encoding='utf-8')
import openpyxl
from docx import Document
B=r"C:\Users\raulh\Desktop\Claude\PlastHub"

# 1) verify corrected formulas
wb=openpyxl.load_workbook(B+r"\PlasHub_Project_Profile_v3_corrected.xlsx",data_only=False)
ws=wb["3b. Flujo de Caja"]
print("R14 D,E,N:", ws.cell(14,4).value,"||",ws.cell(14,5).value,"||",ws.cell(14,14).value)

# 2) recompute corrected cash flow manually
vol=[80]*12
buy=[190]+[130]*11; fle=[75]+[12]*11
rev=[32000]+[19840]*11
ciclo=[45,45,45,40,40,35,35,35,30,30,30,30]
opex=8175; capex=20858
cob=[]
for m in range(12):
    if m==0: cob.append(rev[0] if ciclo[0]<=30 else 0)
    else: cob.append(rev[m] if ciclo[m]<=30 else rev[m-1])
pag=[-vol[m]*(buy[m]+fle[m]) for m in range(12)]
net=[cob[m]+pag[m]-opex for m in range(12)]
caja=[];
for m in range(12):
    base=(-capex if m==0 else 0)+net[m]
    caja.append((caja[m-1] if m>0 else 0)+base)
print("Corrected Caja minima:", min(caja), "(mes", caja.index(min(caja))+1,")")
print("Corrected Caja M12:", round(caja[11]))
print("Pagos M2 (corrected):", pag[1], " vs buggy -21200")

# 3) docx + xlsx structure
for f in ["PlasHub_Business_Plan.docx","MOU_Suministro_Alex_PlasHub_BORRADOR.docx"]:
    d=Document(B+"\\"+f); print(f,"-> paras",len(d.paragraphs),"tables",len(d.tables))
wt=openpyxl.load_workbook(B+r"\PlasHub_Permits_Legal_Tracker.xlsx")
ws2=wt.active; print("Tracker rows:",ws2.max_row,"cols:",ws2.max_column)
