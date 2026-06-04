# -*- coding: utf-8 -*-
"""Build the investor pack: (A) fixed cash-flow xlsx, (B) business-plan docx,
(C) permits/legal tracker xlsx, (D) volume MOU draft docx."""
import io,sys; sys.stdout=io.TextIOWrapper(sys.stdout.buffer,encoding='utf-8')
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE=r"C:\Users\raulh\Desktop\Claude\PlastHub"
CTX=BASE+r"\Context"
DARK=RGBColor(0x25,0x2A,0x31); ORANGE=RGBColor(0xDF,0x69,0x01); TEAL=RGBColor(0x15,0x61,0x6D)
GREY=RGBColor(0x6a,0x72,0x80)

# ============================================================ (A) FIX CASH FLOW
def fix_cashflow():
    src=CTX+r"\PlasHub_Project_Profile_v2.xlsx"
    out=BASE+r"\PlasHub_Project_Profile_v3_corrected.xlsx"
    wb=openpyxl.load_workbook(src,data_only=False)
    ws=wb["3b. Flujo de Caja"]
    # Row 14: fix columns D..N to reference each month's own price+flete column
    for col in range(4,15):           # D=4 ... N=14
        L=get_column_letter(col)
        ws.cell(14,col).value=f"=-{L}6*('2. Supuestos & PL'!{L}7+'2. Supuestos & PL'!{L}9)"
    # annotate the fix
    note=("  [CORREGIDO v3: fila 14 'Pagos a proveedores' ahora usa el precio/flete del mes "
          "correspondiente; antes referenciaba la columna M1 para todos los meses, lo que "
          "inflaba el pago a -$21,200/mes y mostraba insolvencia falsa.]")
    ws.cell(2,2).value=(ws.cell(2,2).value or "")+note
    wb.save(out)
    # report recomputed-ish sanity (values not recalced by openpyxl; Excel will on open)
    print("A) saved",out)

# ============================================================ docx helpers
def newdoc():
    d=Document(); d.styles['Normal'].font.name='Calibri'; d.styles['Normal'].font.size=Pt(10.5); return d
def H(d,t,lvl=1,color=TEAL):
    p=d.add_heading(t,level=lvl)
    for r in p.runs: r.font.color.rgb=color
    return p
def P(d,t,bold=False,italic=False,size=10.5,color=None,after=4):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(after); return p
def B(d,t):
    p=d.add_paragraph(t,style='List Bullet')
    for r in p.runs: r.font.size=Pt(10.5)
    return p
def shade_header(cell,fill='252A31'):
    from docx.oxml.ns import qn; from docx.oxml import OxmlElement
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def TBL(d,headers,rows,widths=None,fontsz=9.5):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; rr=c.paragraphs[0].add_run(h); rr.bold=True; rr.font.size=Pt(fontsz); rr.font.color.rgb=RGBColor(255,255,255); shade_header(c)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=''; rr=cells[i].paragraphs[0].add_run(str(v)); rr.font.size=Pt(fontsz)
            if i>0 and isinstance(v,str) and (v.startswith('$') or v.endswith('%') or v.endswith('t/mo') or v.startswith('-$') or v.startswith('M')):
                cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    d.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

# ============================================================ (B) BUSINESS PLAN
def business_plan():
    d=newdoc()
    tp=d.add_paragraph(); r=tp.add_run('BUSINESS PLAN — Stage 1 (12 meses)'); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=DARK
    r=d.add_paragraph().add_run('PlasHub — Hub de Consolidación de Plásticos · Neo Amaru Global S.A.C. · Factory "Amaru I" · Lima Norte, Perú')
    r.italic=True; r.font.size=Pt(11.5); r.font.color.rgb=TEAL
    P(d,'Documento para inversores / prestamistas · Cifras en USD (€1≈$1.07) · Modelo: comprar → clasificar → pesar → almacenar → consolidar 20 t → despachar contenedor 40′. Sin procesamiento en Stage 1.',italic=True,size=9.5,color=GREY)

    H(d,'1. Resumen ejecutivo',2)
    P(d,'PlasHub convierte una oferta fragmentada y dispersa de plásticos reciclables post-proceso en lotes consolidados, trazables y de grado industrial para exportación. Compra a precios fragmentados, segrega y controla calidad por polímero, acumula ~20 t de un mismo material y despacha un contenedor de 40 pies — reduciendo el tiempo de respuesta del cliente ancla de ~28 a <14 días y capturando el spread entre el precio de compra fragmentado y el de venta consolidado (margen bruto ponderado ~49%, ≈$127/t).')
    TBL(d,['Métrica clave','Valor'],[
        ['Margen bruto ponderado','~49% (≈$127/t) · compra $132/t · venta $259/t'],
        ['Punto de equilibrio','~82 t/mes (Lean A) · ~144 t/mes (Full B)'],
        ['Caso Base (100 t/mes, Lean A)','EBITDA ≈ +$2.0K/mes · equilibrio operativo M3 · ROI ~69%/año sobre CAPEX'],
        ['Capital confirmado','~$22,000 (Alex 60% / Raúl 40%)'],
        ['Necesidad de capital','Lean ≈ $27K CAPEX+WC; con ciclo de caja, pico de financiamiento ≈ $45–50K'],
        ['Objetivo Stage 1','Validar el modelo comercial en 12 meses antes de invertir en procesamiento'],
    ],widths=[2.4,4.1])
    P(d,'El pedido: cerrar la brecha de financiamiento de Stage 1 para una validación de 12 meses. El verdadero cuello de botella no es el CAPEX sino el capital de trabajo: el ciclo "pago ahora / cobro después" lleva el saldo de caja a un pozo de ~−$26K hacia el M4 incluso en el caso Base rentable.',bold=False)

    H(d,'2. El negocio y la propuesta de valor',2)
    for x in ['Disponibilidad garantizada y calidad consistente (ficha técnica + trazabilidad de lote).',
              'Tiempo de respuesta 28 → <14 días; lotes repetidos ~2 semanas según stock.',
              'Captura del spread compra-vs-consolidado que Alex hoy pierde por urgencia.',
              'Eslabón físico que "amarra" a Alex como cliente de largo plazo y es replicable en LatAm.']:
        B(d,x)
    P(d,'Tesis por etapas: Stage 1 (este plan) consolidación ~80–100 t/mes, ~900 m²; Stage 2 procesamiento (compactado/lavado/peletizado) que desbloquea el tramo de precio "regrind" $900–1,300/MT, 400–500+ t/mes; Stage 3 réplica regional (~6,000 m²).',italic=True,size=9.5,color=GREY)

    H(d,'3. Materiales y precios de mercado (Perú/LatAm 2025–26)',2)
    P(d,'Tramo "feedstock embalado → lote consolidado" — el producto que PlasHub realmente comercializa. Triangulado de: referencia Lima-2026 de la empresa, ancla verificada Polvillo de Flake ($150 compra/$300 venta), y rangos de acopio en Perú. Nota: las cifras "HDPE reciclado $930–1,310/MT" muy citadas son regrind lavado/peletizado — un producto distinto y de mayor valor (la oportunidad de Stage 2), no el feedstock de Stage 1.')
    TBL(d,['Material','Frec.','Mix %','Compra $/t','Venta $/t','Margen $/t','Margen %'],[
        ['PET post-consumo','Frec.','27%','$120','$240','$120','50%'],
        ['HDPE industrial','Frec.','26%','$140','$270','$130','48%'],
        ['PP post-industrial','Frec.','21%','$130','$260','$130','50%'],
        ['PE Film','Frec.','16%','$100','$200','$100','50%'],
        ['ABS / PS industrial','Espor.','4%','$200','$380','$180','47%'],
        ['Nylon / PA','Espor.','3%','$280','$520','$240','46%'],
        ['Special mix','Espor.','3%','$120','$230','$110','48%'],
        ['Polvillo de Flake (ancla)','Frec.','—','$150','$300','$150','50%'],
        ['PROMEDIO PONDERADO','—','100%','$132','$259','$127','49%'],
    ],widths=[1.9,0.7,0.6,0.85,0.85,0.85,0.7],fontsz=9)
    P(d,'Honestidad: no existen índices públicos de scrap embalado en Perú; compra/venta son cifras de referencia trianguladas — validar con 2–3 cotizaciones reales de proveedores antes de comprometer capital. Acopio Perú: PET S/0.70–2.60/kg; HDPE/PP/LDPE S/0.30–0.80/kg.',italic=True,size=9,color=GREY)

    H(d,'4. Modelo financiero — tres escenarios',2)
    P(d,'Supuestos: rampa M1 50% / M2 75% / M3 90% / M4–12 100%; montacarga alquilado (OPEX) en ambos setups; cobros con desfase ~1 mes (clientes pagan a 30–45 d), pagos a proveedores contra entrega — el ciclo de caja adverso está modelado explícitamente. Costo variable (flete inbound + drayage a puerto) ≈ $15/t.')
    TBL(d,['Métrica (estado estable M4+)','Conservador 40 t · A','Base 100 t · A','Optimista 200 t · B'],[
        ['Ingresos /mes','$10,372','$25,930','$51,860'],
        ['Utilidad bruta /mes','$5,088','$12,720','$25,440'],
        ['OPEX fijo /mes','$9,225','$9,225','$16,120'],
        ['EBITDA /mes','-$4,737','+$1,995','+$6,320'],
        ['Equilibrio (volumen)','82 t/mo','82 t/mo','144 t/mo'],
        ['Equilibrio operativo','nunca','M3','M2'],
        ['Pozo de caja (mín.)','-$69,889','-$26,345','-$64,788'],
        ['EBITDA año 1','-$53,800','+$14,403','+$60,900'],
        ['ROI sobre CAPEX @M12','-291%','+69%','+142%'],
    ],widths=[2.3,1.4,1.3,1.4],fontsz=9)
    P(d,'Lectura: el caso Conservador (40 t/mes) es deficitario — está por debajo del equilibrio (82 t/mes) y quema caja; es un escenario de alerta, no viable como negocio autónomo. El caso Base es rentable desde el M3 pero con margen neto delgado. El Optimista es atractivo (ROI 142%) pero requiere el setup Full (~$50K). En los tres casos el pozo de caja exige financiar el capital de trabajo, no solo el CAPEX.')

    H(d,'5. CAPEX — itemizado, ambos setups',2)
    H(d,'5.1 Scenario A — Lean (subarriendo ~900 m²)',3,DARK)
    TBL(d,['Línea','USD','Nota'],[
        ['Garantía nave (2 meses)','$6,372','Negociable a 1 mes → −$3,186'],
        ['Adelanto 1 mes de renta','$3,186',''],
        ['Adecuaciones: demarcación 7 zonas, señalética INDECI, luminarias, pintura piso','$1,500','Floor marking & signage'],
        ['Báscula de piso 3 t','$1,200','Compra'],
        ['Transpaleta + EPP + herramientas','$1,100','Montacarga alquilado (OPEX)'],
        ['Consultor EC-RS / ambiental (MINAM)','$3,500','Ruta crítica 60–90 d'],
        ['Asesoría legal (MOU, pacto socios, licencia func.)','$1,500',''],
        ['Ingeniería / diseño de layout','$1,000',''],
        ['Registro SUNAT (RMT) + contabilidad inicial','$1,500',''],
        ['CAPEX subtotal','$20,858',''],
        ['+ Capital de trabajo (ciclo de caja)','$6,000','financia el pozo M1–3'],
        ['TOTAL necesidad de capital','$26,858','vs confirmado $22K'],
    ],widths=[3.4,0.9,2.0],fontsz=9)
    H(d,'5.2 Scenario B — Full (lease directo 1,500 m²)',3,DARK)
    TBL(d,['Línea','USD','Nota'],[
        ['Garantía nave (2 meses @ $7,500)','$15,000','Factor swing — a 1 mes ahorra $7,500'],
        ['Adelanto 1 mes de renta','$7,500',''],
        ['Fit-out / obra civil (oficina, racks, piso, demarcación)','$4,500','incl. floor marking & signage'],
        ['Báscula de piso 3 t','$1,500',''],
        ['Racking esporádicos + EPP + herramientas','$2,500',''],
        ['Consultor EC-RS / ambiental','$3,500',''],
        ['Asesoría legal (pacto socios, licencias)','$2,000',''],
        ['Ingeniería / diseño layout','$1,200',''],
        ['Registro SUNAT (RMT) + contabilidad inicial','$1,500',''],
        ['Sistema de inventario (software)','$800',''],
        ['CAPEX subtotal','$40,000',''],
        ['+ Capital de trabajo','$10,000',''],
        ['TOTAL necesidad de capital','$50,000','vs target $50K'],
    ],widths=[3.4,0.9,2.0],fontsz=9)

    H(d,'6. OPEX mensual fijo',2)
    TBL(d,['Concepto','Lean A $/mo','Full B $/mo'],[
        ['Renta nave','$3,186 (~900 m²)','$7,500 (1,500 m²)'],
        ['Montacarga 3 t (alquiler)','$1,000','$1,000'],
        ['Jefe de Planta / Site Manager','$1,100','$1,300'],
        ['Operarios de planta','$1,100 (×2)','$1,650 (×3)'],
        ['Supervisor / QA (parcial)','—','$600'],
        ['Electricidad trifásica','$350','$600'],
        ['Seguridad / vigilancia','$1,100','$1,200'],
        ['Mantenimiento','$50','$150'],
        ['Contabilidad / contador','$500','$650'],
        ['EPP / consumibles','$200','$350'],
        ['Comunicaciones','$100','$150'],
        ['Servicios básicos','$100','$200'],
        ['Contingencia 5%','$439','$770'],
        ['TOTAL fijo /mes','$9,225','$16,120'],
    ],widths=[2.6,1.9,1.9],fontsz=9)
    P(d,'Costo variable ≈ $15/t (flete inbound + drayage a puerto), modelado aparte. Salarios reflejan RMV Perú 2026 (S/1,130 ≈ $305) + beneficios (CTS, gratificaciones, EsSalud). Raúl no cobra salario hasta EBITDA ≥ 0 por 2 meses consecutivos.',italic=True,size=9,color=GREY)

    H(d,'7. Brecha de capital y plan de financiamiento',2)
    P(d,'Hay dos brechas, no una: (1) la brecha de activos (CAPEX) — equipos + garantía + permisos; y (2) la brecha de capital de trabajo — el ciclo compra-ahora/cobra-después lleva el saldo de banco a negativo en M1–M3 aunque el P&L sea positivo. Esta segunda es la mayor y la razón #1 por la que pilotos como este se estancan.')
    P(d,'Palancas de financiamiento (ranking por apalancamiento/costo):')
    for x in ['1. Negociar entrada (1 mes de garantía vs 2–3) → ahorra $3–15K. La palanca más alta y barata.',
              '2. Comprimir el ciclo de caja — Alex prepaga / adelanto parcial, o venta a plazo corto / cobro contra entrega.',
              '3. Préstamo entre socios (notariado) para complementar los $22K confirmados.',
              '4. Faseo del CAPEX — escalonar permisos / fit-out en M1–M3.',
              '5. Línea de capital de trabajo MYPE (BCP/Interbank/Mibanco, TEA ~25–45% — caro; dimensionar solo al pozo) o factoring/confirming sobre cuentas por cobrar.',
              '6. Diferir salario de Raúl hasta EBITDA positivo (ya en el modelo).']:
        B(d,x)
    P(d,'Recomendación: iniciar Lean (A), cerrar el pozo de ~$15–18K con negociación de garantía + prepago de Alex + un pequeño puente entre socios, validar el caso Base (100 t/mes), y luego graduar a Full (B) y Stage 2 financiados con utilidades retenidas + una línea bancaria con volumen ya probado.',bold=True)

    H(d,'8. Registro de riesgos',2)
    TBL(d,['Riesgo','Prob.','Impacto','Mitigación'],[
        ['EC-RS no iniciado (MINAM)','Alta','Crítico — bloquea operación legal','Consultor ambiental día 1; abrir SIGERSOL; Plan de Manejo de Residuos. 60–90 d.'],
        ['Capital de trabajo / ciclo de caja','Alta','Alto — banco negativo M1–3','Prepago Alex, garantía 1 mes, puente socios, factoring.'],
        ['Sin compromiso de volumen (Alex)','Media','Alto — modelo inválido','Carta/contrato: mín. t/mes por 12 meses antes de firmar nave.'],
        ['Volumen bajo equilibrio (~82 t/mes)','Media','Alto — deficitario','Reducir OPEX (renta) antes de buscar volumen. Alerta KPI <60% target.'],
        ['Concentración en cliente único','Media','Alto','Sumar 1–2 clientes; formalizar forecast; diversificar en Stage 2.'],
        ['Precio de transferencia (SUNAT)','Media','Alto — Alex es proveedor+cliente+socio','Facturas a precio de mercado, documentadas; registrar en pacto de socios.'],
        ['Retraso de permisos (EIA/ITSE)','Alta','Crítico (retrasa SOP)','Asesor experto, trámites en paralelo, buffer 6–12 meses.'],
        ['Compresión de precios (venta)','Media','Medio','Estresado en el modelo; diversificar; migrar a tramo regrind (Stage 2).'],
    ],widths=[1.9,0.6,1.7,2.3],fontsz=9)

    H(d,'9. Hoja de ruta 12 meses y Gate',2)
    TBL(d,['Fase','Meses','Hitos'],[
        ['Setup & permisos','M0–M3','EC-RS (inicio día 1); pacto de socios + RMT; nave firmada; layout; montacarga + báscula; primer inbound.'],
        ['Rampa','M1–M3','Volumen 50→90%; base de proveedores + QC; primeros contenedores consolidados.'],
        ['Operación estable','M4–M9','~100 t/mes (Base); KPIs: rechazo <10%, margen ≥40%, días inventario <30, caja ≥$3K.'],
        ['Validar & Gate','M10–M12','Probar margen y fiabilidad; Gate 1: Go / cambiar alcance / parar inversión Stage 2.'],
    ],widths=[1.4,0.9,4.2],fontsz=9)

    H(d,'10. Fuentes y supuestos',2)
    P(d,'Precios regrind (benchmark): ChemAnalyst, IMARC, Intratec. Feedstock Perú: QuimiNet + acopio (PET S/0.70–2.60/kg; HDPE/PP/LDPE S/0.30–0.80/kg). Renta Lima Norte: Adondevivir, Nuroa ($2.75–9.5/m²+IGV). RMV 2026 S/1,130: Infobae, BCRP. Datos de empresa: PlasHub_Project_Profile_v2.xlsx, Requerimientos_Almacen_Peru.pptx.',size=9,color=GREY)
    P(d,'Notas de honestidad: (1) compra/venta son referencias trianguladas — validar con cotizaciones reales. (2) La pestaña de flujo de caja del Profile v2 tenía un error de fórmula (pagos fijados al precio M1 todo el año) corregido en PlasHub_Project_Profile_v3_corrected.xlsx. (3) "Margen %" es bruto (venta−compra)/venta, antes de OPEX; el margen neto por tonelada en el caso Base es delgado (~$20/t tras todos los costos).',italic=True,size=9,color=GREY)

    out=BASE+r"\PlasHub_Business_Plan.docx"; d.save(out); print("B) saved",out)

# ============================================================ (C) PERMITS TRACKER
def permits_tracker():
    wb=openpyxl.Workbook(); ws=wb.active; ws.title="Permisos & Legal"
    title=ws.cell(1,1,"PlasHub / Factory \"Amaru I\" — Ruta crítica de Permisos y Legal (Stage 1)")
    title.font=Font(bold=True,size=14,color="252A31"); ws.merge_cells('A1:L1')
    sub=ws.cell(2,1,"Iniciar HOY los ítems de ruta crítica (✦). Duraciones en días hábiles aprox. Actualizar 'Estado' semanalmente.")
    sub.font=Font(italic=True,size=10,color="6A7280"); ws.merge_cells('A2:L2')
    headers=["#","Ítem","Categoría","Responsable","Duración (d)","Depende de","Ruta crítica","Inicio objetivo","Fin objetivo","Costo (USD)","Estado","Notas"]
    hr=4
    for i,h in enumerate(headers,1):
        c=ws.cell(hr,i,h); c.font=Font(bold=True,color="FFFFFF",size=9.5); c.fill=PatternFill("solid",fgColor="252A31")
        c.alignment=Alignment(horizontal="center",vertical="center",wrap_text=True)
    rows=[
        [1,"EC-RS — registro Empresa Comercializadora de Residuos Sólidos (MINAM/SIGERSOL)","Ambiental","Asesor ambiental","60–90","—","✦ SÍ","","","3,500","No iniciado","BLOQUEA operación legal. Abrir cuenta SIGERSOL + Plan de Manejo de Residuos. INICIAR DÍA 1."],
        [2,"Constitución / verificación SAC + cambio a Régimen MYPE Tributario (RMT)","Tributario","Contador","10–15","—","✦ SÍ","","","1,000","En progreso","Ventas año 1 > límite RER → RMT obligatorio ANTES de la 1ra venta."],
        [3,"Pacto de socios (notarial): equity 60/40, precio transferencia, dividendos","Legal/Societario","Abogado + socios","10–20","—","✦ SÍ","","","1,000","No iniciado","Firmar ANTES de la primera operación. Mitiga riesgo SUNAT partes vinculadas."],
        [4,"Carta/contrato de compromiso de volumen mínimo (Alex)","Comercial","Raúl + Alex","7–14","—","✦ SÍ","","","0","No iniciado","Mín. t/mes por 12 meses. Sin esto el modelo financiero es inválido."],
        [5,"Selección y firma de nave (subarriendo ~900 m²)","Inmobiliario","Raúl + Alex","20–30","1,4","✦ SÍ","","","9,558","No iniciado","No firmar antes de cerrar brecha de capital. Negociar garantía 1 mes."],
        [6,"Compatibilidad de uso / zonificación (municipio)","Municipal","Asesor legal","15–30","5","SÍ","","","200","No iniciado","Verificar que el predio permite almacenar plásticos (residuo no peligroso)."],
        [7,"Licencia de funcionamiento (municipal)","Municipal","Asesor legal","15–45","5,6","SÍ","","","500","No iniciado","Depende de zonificación + ITSE."],
        [8,"ITSE — Inspección Técnica de Seguridad en Edificaciones (INDECI)","Seguridad","Asesor + Jefe Planta","15–30","5","SÍ","","","800","No iniciado","Extintores, salidas, señalética. Ver checklist INDECI en plashub-layout.html."],
        [9,"EIA / DIA (si la autoridad lo exige para Stage 1)","Ambiental","Asesor ambiental","30–180","1","Posible","","","0","Por confirmar","Confirmar si aplica a consolidación (sin procesamiento). Buffer 6–12 m."],
        [10,"Adecuación nave: demarcación 7 zonas, señalética, luminarias, piso","Obra","Raúl + contratista","10–15","5","No","","","1,500","No iniciado","Floor marking & signage."],
        [11,"Montacarga 3 t — contrato de alquiler","Equipos","Raúl","5–10","5","✦ SÍ","","","1,000/mo","No iniciado","Obligatorio para bolsones de 1 t. NO usar transpaletas para apilar."],
        [12,"Báscula de piso 3 t — compra + calibración","Equipos","Raúl","5–10","5","No","","","1,200","No iniciado","Pesaje al ingreso/egreso. Calibración anual."],
        [13,"Seguro integral (responsabilidad, incendio, contenido)","Riesgo","Asesor","7–14","5","No","","","incl. OPEX","No iniciado","Carga de fuego alta (plásticos) — confirmar cobertura."],
        [14,"Cierre de brecha de capital / capital de trabajo","Financiero","Raúl + Alex","15–30","4","✦ SÍ","","","—","No iniciado","Garantía 1 mes + prepago Alex + puente socios. Ver §7 business plan."],
        [15,"Cuenta bancaria empresa + (opcional) línea MYPE / factoring","Financiero","Raúl","10–20","2","No","","","—","No iniciado","Dimensionar línea al pozo de caja (~$26K)."],
        [16,"SOP — inicio de operaciones (primer inbound)","Operación","Jefe de Planta","—","1,3,5,8,11","Hito","","","—","No iniciado","Solo tras EC-RS + ITSE + nave + equipos."],
    ]
    fills={"No iniciado":"FBE9E7","En progreso":"FDF4E3","Completado":"E8F5EE","Por confirmar":"EEF1F4"}
    for r,row in enumerate(rows,hr+1):
        for c,val in enumerate(row,1):
            cell=ws.cell(r,c,val); cell.font=Font(size=9.5); cell.alignment=Alignment(vertical="top",wrap_text=True)
            if c==7 and "SÍ" in str(val): cell.font=Font(size=9.5,bold=True,color="C0392B")
            if c==11: cell.fill=PatternFill("solid",fgColor=fills.get(val,"FFFFFF"))
    widths=[4,34,15,16,11,11,11,13,13,11,13,40]
    for i,w in enumerate(widths,1): ws.column_dimensions[get_column_letter(i)].width=w
    ws.freeze_panes="A5"
    # legend
    lr=hr+len(rows)+2
    ws.cell(lr,2,"Leyenda Estado:").font=Font(bold=True,size=9.5)
    for j,(k,f) in enumerate(fills.items()):
        c=ws.cell(lr,3+j,k); c.fill=PatternFill("solid",fgColor=f); c.font=Font(size=9)
    ws.cell(lr+1,2,"✦ = ruta crítica (determina la fecha de SOP). Iniciar EC-RS, RMT, pacto de socios y compromiso de Alex de inmediato.").font=Font(italic=True,size=9,color="6A7280")
    out=BASE+r"\PlasHub_Permits_Legal_Tracker.xlsx"; wb.save(out); print("C) saved",out)

# ============================================================ (D) MOU / TERM SHEET
def mou_draft():
    d=newdoc()
    r=d.add_paragraph().add_run('BORRADOR — NO FIRMAR SIN REVISIÓN LEGAL'); r.bold=True; r.font.color.rgb=RGBColor(0xC0,0x39,0x2B); r.font.size=Pt(10)
    tp=d.add_paragraph(); r=tp.add_run('MEMORÁNDUM DE ENTENDIMIENTO (MOU) Y TÉRMINOS DE SUMINISTRO'); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=DARK
    r=d.add_paragraph().add_run('Acuerdo de Suministro Mínimo, Comercial y de Gobernanza · PlasHub / Neo Amaru Global S.A.C.')
    r.italic=True; r.font.size=Pt(11); r.font.color.rgb=TEAL
    P(d,'Borrador de trabajo para revisión de un abogado/notario peruano. No constituye asesoría legal. Los espacios "____" deben completarse al firmar.',italic=True,size=9,color=GREY)

    H(d,'1. Partes',2)
    P(d,'1.1. ____________________ ("Alex"), identificado con ____, en calidad de socio, proveedor de materiales y cliente ancla.')
    P(d,'1.2. Neo Amaru Global S.A.C. ("PlasHub"), RUC ____, representada por Raúl ____________________, en calidad de operador del hub de consolidación.')

    H(d,'2. Antecedentes',2)
    P(d,'PlasHub opera un hub de recepción, clasificación, pesaje, almacenamiento y consolidación de plásticos post-proceso en Lima Norte (Stage 1, sin procesamiento). Alex aporta flujo comercial, red de proveedores y demanda. Este MOU define el compromiso mínimo de volumen, los precios, los estándares de calidad y la estructura societaria necesarios para sustentar la inversión de Stage 1.')

    H(d,'3. Objeto',2)
    P(d,'Establecer (i) un compromiso de volumen mínimo que haga viable el modelo financiero de Stage 1, (ii) el mecanismo de precios entre las partes, y (iii) la gobernanza y distribución societaria.')

    H(d,'4. Compromiso de volumen mínimo',2)
    P(d,'4.1. Alex se compromete a canalizar a PlasHub un volumen mínimo de ______ t/mes (referencia: ≥82 t/mes para superar el punto de equilibrio; objetivo 100 t/mes) durante 12 meses, a partir del mes ____.')
    P(d,'4.2. Si el volumen real cae por debajo de ____ t/mes por dos meses consecutivos, las partes convocarán una reunión de revisión en un plazo de ____ días.')
    P(d,'4.3. Mix referencial de materiales: PET, HDPE, PP, PE Film (frecuentes) y ABS/PS, Nylon/PA, mix especiales (esporádicos).')

    H(d,'5. Precios y margen',2)
    P(d,'5.1. PlasHub compra a precio de mercado verificable y vende el lote consolidado al cliente final (Modelo A — PlasHub comercializa y captura el margen).')
    P(d,'5.2. Precios de referencia (USD/t, ajustables por acta): PET 120/240 · HDPE 140/270 · PP 130/260 · PE Film 100/200 · ABS-PS 200/380 · Nylon/PA 280/520 · Flake 150/300 (compra/venta).')
    P(d,'5.3. Precio de transferencia: toda transacción entre partes vinculadas se documentará a precio de mercado, con facturas de respaldo, para cumplimiento ante SUNAT. NUNCA se usarán precios artificiales.')

    H(d,'6. Calidad, trazabilidad y rechazos',2)
    P(d,'6.1. Cada bolsón se pesa y clasifica al ingreso; material contaminado o fuera de especificación se rechaza (objetivo de rechazo <10% en Stage 1, <5% en Stage 2).')
    P(d,'6.2. PlasHub mantiene trazabilidad por lote (libro EC-RS + báscula) y emite ficha técnica del lote consolidado.')

    H(d,'7. Estructura societaria y gobernanza',2)
    P(d,'7.1. Participación: Alex 60% / Raúl 40%, a formalizar en el Estatuto SAC ante SUNARP antes de la primera operación.')
    P(d,'7.2. Salario de Raúl: inicia (S/3,500–4,500/mes) recién tras EBITDA ≥ 0 por dos meses consecutivos.')
    P(d,'7.3. Dividendos: no se distribuyen hasta recuperar el CAPEX (EBITDA acumulado > CAPEX). Reserva del 30% de utilidades para reinversión.')
    P(d,'7.4. Aportes de capital: Alex ~$16,000; Raúl ~$5,500 (ago-2026). Cualquier puente entre socios se documentará por separado (préstamo notariado).')
    P(d,'7.5. Decisiones que requieren acuerdo conjunto: CAPEX > $____, cambios de alcance, endeudamiento, incorporación de nuevos socios/clientes ancla.')

    H(d,'8. Plazo, confidencialidad y naturaleza',2)
    P(d,'8.1. Vigencia: 12 meses (Stage 1), renovable por acuerdo escrito. Gate 1 al M12 decide la continuidad hacia Stage 2.')
    P(d,'8.2. Confidencialidad: precios, proveedores y clientes son información reservada de PlasHub.')
    P(d,'8.3. Naturaleza: las cláusulas 4, 5 y 7 son vinculantes una vez constituida la SAC y firmado el pacto de socios; el resto expresa la intención de las partes y se perfeccionará en los contratos definitivos.')

    H(d,'9. Firmas',2)
    TBL(d,['Parte','Nombre','Firma','Fecha'],[
        ['Socio / Proveedor / Cliente ancla','Alex','_____________________','__________'],
        ['Operador (PlasHub / Neo Amaru Global S.A.C.)','Raúl','_____________________','__________'],
        ['Testigo / Asesor legal','—','_____________________','__________'],
    ],widths=[2.6,1.4,1.6,1.0],fontsz=9.5)
    out=BASE+r"\MOU_Suministro_Alex_PlasHub_BORRADOR.docx"; d.save(out); print("D) saved",out)

fix_cashflow()
business_plan()
permits_tracker()
mou_draft()
print("ALL DONE")
