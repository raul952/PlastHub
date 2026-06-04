# -*- coding: utf-8 -*-
"""Build the corrected Project Charter v2.0 with validated Stage-1 numbers."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DARK=RGBColor(0x25,0x2A,0x31); ORANGE=RGBColor(0xDF,0x69,0x01); TEAL=RGBColor(0x15,0x61,0x6D)
OUT=r"C:\Users\raulh\Desktop\Claude\PlastHub\Amaru-I-Project-Charter-v2.docx"

doc=Document()
st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)

def h(text,level=1,color=TEAL):
    p=doc.add_heading(text,level=level)
    for r in p.runs: r.font.color.rgb=color
    return p
def para(text,bold=False,italic=False,size=10.5,color=None,after=4):
    p=doc.add_paragraph(); r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(after); return p
def bullet(text):
    p=doc.add_paragraph(text,style='List Bullet');
    for r in p.runs: r.font.size=Pt(10.5)
    return p
def table(headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=t.rows[0].cells
    for i,htxt in enumerate(headers):
        hdr[i].text=''; rr=hdr[i].paragraphs[0].add_run(htxt); rr.bold=True; rr.font.size=Pt(9.5); rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        from docx.oxml.ns import qn; from docx.oxml import OxmlElement
        tcPr=hdr[i]._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'252A31'); tcPr.append(shd)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=''; rr=cells[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(9.5)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

# ---- TITLE ----
tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.LEFT
r=tp.add_run('PROJECT CHARTER — v2.0 (Validado)'); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=DARK
sp=doc.add_paragraph(); r=sp.add_run('Factory "Amaru I" / PlasHub — Planta Piloto de Recolección, Clasificación y Almacenamiento de Plásticos Post-Proceso · Perú')
r.font.size=Pt(11.5); r.font.color.rgb=TEAL; r.italic=True

bn=doc.add_paragraph(); bn.paragraph_format.space_before=Pt(6)
r=bn.add_run('⚠ Versión 2.0 — actualizada con las cifras validadas del Project Profile v2, la presentación de requerimientos del almacén y el documento de proyecto piloto. '
             'Reemplaza la v1.0 (02-feb-2026), cuyas cifras (500 t/mes, 2,000–3,000 m², CAPEX USD 5–15K) correspondían a etapas posteriores (Stage 2/3), no a la operación piloto Stage 1.')
r.font.size=Pt(9.5); r.italic=True; r.font.color.rgb=ORANGE

# ---- CORRECTION LOG ----
h('Registro de correcciones v1.0 → v2.0',2,ORANGE)
table(['Concepto','v1.0 (desactualizado)','v2.0 (validado)','Fuente'],
[['Volumen mensual Stage 1','500 ton/mes almacenadas','80 t/mes (base) · 100 t/mes (optimista) · 61 t/mes break-even','Project Profile v2 — Supuestos & P&L'],
 ['Volumen semanal','100 ton/sem (20 t × 5 cargas)','~10 t/sem · 1–2 camiones/sem · 8–12 bolsones','Requerimientos Almacén (PPTX)'],
 ['Bolsón / big bag','800×1200 mm, 600–1000 kg','1 m³ (1×1×1 m), ~1 t, apilado máx. 2 alto','PPTX / Profile'],
 ['Área nave Stage 1','2,000–3,000 m²','~900 m² (útil 500–750; hasta 1,000). 6,000 m² = Stage 3','Profile (renta) / PPTX'],
 ['CAPEX Stage 1','USD 5–15K (tabla: 5–50K)','USD 20,858 (capital disp. $20–22K; gap −$858)','Project Profile v2 — CAPEX'],
 ['OPEX mensual','no especificado','USD 8,175/mes (renta $3,186 ≈ 900 m²)','Project Profile v2 — OPEX'],
 ['Margen objetivo','"≥20%" / ">50%"','Bruto pond. ~49% ($128.7/t); neto ~$3.8/t a 80 t/mes','Profile — Margen por material'],
 ['Ubicación','Huachipa / Chancay','Lima Norte (SMP / Los Olivos / Puente Piedra)','Profile — Project Background'],
 ['Equipo de manejo','"montacargas, racks"','Montacarga contrabalanceado 3 t (obligatorio) + báscula 3 t','Profile / PPTX / Roles'],
])

# ---- 1 INFO ----
h('1. Información general del proyecto',2)
table(['Ítem','Detalle'],
[['Nombre del proyecto','Factory "Amaru I" / PlasHub — Planta Piloto (Stage 1)'],
 ['Empresa','Neo Amaru Global S.A.C.'],
 ['Lugar','Perú — Lima Norte (SMP / Los Olivos / Puente Piedra)'],
 ['Tipo de proyecto','Hub de recepción, clasificación, pesaje, almacenamiento y consolidación (sin transformación)'],
 ['Etapa actual','Stage 1 — Piloto de consolidación (sin trituración / lavado / peletizado)'],
 ['Fecha de inicio','Febrero 2026'],
 ['Duración estimada','Stage 1: operación estable y validación 6–9 meses'],
 ['CAPEX autorizado (Stage 1)','USD 20,858'],
 ['Capital disponible','USD 20–22K (Alex ~$15K + Raúl ~$5K, desde ago-2026) · gap −$858 a cerrar'],
 ['Modelo comercial','Modelo A — PlasHub compra y vende (captura margen completo)'],
 ['Estructura societaria','Alex 60% (sponsor / proveedor / cliente ancla) · Raúl 40% (PM / operación)'],
 ['Project Manager','Raúl'],
 ['Patrocinador ejecutivo','Alex'],
 ['Versión / Fecha','2.0 · (actualiza v1.0 del 02-feb-2026)'],
],widths=[2.2,4.3])

# ---- 2 PURPOSE ----
h('2. Propósito y justificación',2)
h('2.1 Problema identificado',3,DARK)
para('Alex, trader de commodities, debe consolidar volumen de plásticos post-proceso a partir de una red dispersa de proveedores pequeños. Esto genera: disponibilidad fragmentada, pérdida de margen por urgencia, variabilidad de calidad (sin ficha técnica ni trazabilidad) y falta de un eslabón físico propio que controle calidad y logística de consolidación.')
h('2.2 Solución propuesta',3,DARK)
para('PlasHub actúa como "pulmón de inventario": RECIBE → CLASIFICA → PESA → ALMACENA → CONSOLIDA ~20 t por material → DESPACHA un contenedor de 40 pies. El proyecto se ejecuta por etapas, dimensionando la inversión a resultados:')
bullet('Stage 1 — Consolidación pura (este charter): recolección, clasificación, pesaje y almacenamiento segregado. ~80 t/mes, nave ~900 m². Sin maquinaria de transformación.')
bullet('Stage 2 — Procesamiento básico: compactado / empaque para estandarizar producto y mejorar densidad de carga. Volumen objetivo 400–500+ t/mes.')
bullet('Stage 3 — Optimización y réplica regional: escala a 5+ países, nave ~6,000 m², procesamiento avanzado.')
h('2.3 Beneficios esperados',3,DARK)
table(['Beneficio','Descripción','Responsable'],
[['Mejora de margen','Reducir pérdida por urgencia; margen bruto ponderado ~49% ($128.7/t) con calidad controlada.','Raúl (KPIs)'],
 ['Disponibilidad','Stock consolidado listo; tiempos de respuesta más cortos vs. búsqueda contra reloj.','Raúl + Alex'],
 ['Calidad consistente','Rechazos < 10% (Stage 1) / < 5% (Stage 2); trazabilidad por lote (libro EC-RS + báscula).','Responsable de Calidad'],
 ['Escalabilidad','Modelo documentado y replicable en LatAm.','Raúl'],
 ['Relación socio','Plataforma estable para que Alex concentre energía en ventas.','Alex + Raúl'],
])

# ---- 3 OBJECTIVES ----
h('3. Objetivos del proyecto',2)
h('3.1 Objetivo primario',3,DARK)
para('Implementar y operar un hub de consolidación de plásticos post-proceso en Lima Norte que garantice disponibilidad de inventario, control de calidad y márgenes mejorados para Alex y sus clientes, demostrando la viabilidad del modelo para réplica regional.')
h('3.2 Objetivos específicos — Stage 1 (Piloto / Almacenaje)',3,DARK)
bullet('Obtener los permisos requeridos: EC-RS (MINAM/SIGERSOL) — CRÍTICO, 60–90 días, iniciar de inmediato; licencia municipal de funcionamiento; INDECI / ITSE; EIA si la autoridad lo exige.')
bullet('Asegurar y adecuar una nave ~900 m² (útil 500–750 m²) con patio de maniobra ≥200 m², altura libre ≥5 m, piso industrial ≥5 t/m² y 7 zonas segregadas demarcadas.')
bullet('Instalar equipos Stage 1: montacarga contrabalanceado eléctrico de 3 t (obligatorio para bolsones de 1 t), báscula de piso 3 t, EPP y sistema simple de registro/inventario.')
bullet('Lograr operación estable a ~80 t/mes (break-even 61 t/mes; objetivo optimista 100 t/mes), consolidando ~20 t por material por contenedor.')
bullet('Validar el caso de negocio: margen bruto ponderado ~49%; rechazo por calidad < 10%; cumplimiento de KPIs operativos y de liquidez.')
para('Stage 2 (Procesamiento básico): instalar compactado/empaque; procesar 400–500+ t/mes con valor añadido; margen de procesamiento a validar. '
     'Stage 3 (Réplica): documentar procesos y evaluar apertura en otro país.', italic=True, size=9.5, color=RGBColor(0x6a,0x72,0x80))

# ---- 4 SCOPE ----
h('4. Alcance del proyecto',2)
h('4.1 Incluye (Stage 1)',3,DARK)
for x in ['Selección, aseguramiento y adecuación de la nave/layout en Lima Norte.',
          'Gestión y obtención de permisos (EC-RS, municipal, INDECI/ITSE, EIA si aplica).',
          'Adecuaciones básicas: demarcación de 7 zonas, señalética INDECI, luminarias, corredor peatonal.',
          'Adquisición y puesta en marcha de equipos Stage 1 (montacarga 3 t, báscula 3 t, EPP, registro de inventario).',
          'Definición de procesos (recepción, pesaje, clasificación/QC, almacenaje, consolidación, despacho) y KPIs.',
          'Reclutamiento del equipo mínimo (Jefe de Planta/QA/HSE + operador de montacarga).']:
    bullet(x)
h('4.2 Excluye (Stage 1)',3,DARK)
for x in ['Trituración, lavado, peletizado o cualquier transformación (corresponde a Stage 2+).',
          'Actividad comercial / trading de Alex más allá de aportar demanda y materiales.',
          'Distribución final al cliente último (logística de Alex).']:
    bullet(x)
h('4.3 Restricciones y supuestos',3,DARK)
para('Restricciones:',bold=True,after=2)
for x in ['CAPEX Stage 1 ~USD 20,858; capital disponible $20–22K (gap −$858 a cerrar antes de firmar la nave).',
          'Permisos en Perú pueden tomar 6–12 meses; EC-RS es ruta crítica (60–90 días).',
          'Dependencia de una nave adecuada en zona industrial/logística compatible con plásticos.']:
    bullet(x)
para('Supuestos:',bold=True,after=2)
for x in ['Alex aporta volumen mínimo comprometido (objetivo ~80 t/mes, mín. operativo 50–61 t/mes) — requiere carta/contrato.',
          'Existe mercado local con demanda estable y precios de referencia verificables (Lima 2026).',
          'Régimen tributario RMT (ventas año 1 > límite RER); pacto de socios notarial firmado antes de operar.']:
    bullet(x)

# ---- 5 ECONOMICS ----
h('5. Resumen económico Stage 1 (validado)',2)
table(['Métrica','Valor','Nota'],
[['Volumen objetivo','80 t/mes','base · 100 optimista · 61 break-even'],
 ['Margen bruto ponderado','$128.7/t (~49%)','compra $132.6/t · venta $261.3/t'],
 ['OPEX fijo mensual','USD 8,175','renta $3,186 (~900 m²) es el mayor rubro'],
 ['EBITDA mensual','M1 $2,625 · M2+ ~$305','margen neto delgado a 80 t/mes'],
 ['CAPEX inicial','USD 20,858','báscula + montacarga + adecuación + depósito + permisos'],
 ['Payback estimado','8–12 meses','escenario 80 t/mes (15–18 m si <50 t/mes los 1ros 3 meses)'],
])
para('Nota: la hoja de Flujo de Caja (3b) del Profile v2 contiene un error de enlace (pagos a proveedores fijados en −$21,200 los 12 meses en lugar de ~−$11,360 desde M2), '
     'que muestra insolvencia (−$142,958) pese a un P&L positivo (+$305/mes). Corregir antes de usar el modelo para financiamiento.',
     italic=True,size=9,color=ORANGE)

# ---- 6 GOVERNANCE ----
h('6. Gobernanza y responsabilidades',2)
table(['Rol','Responsable','Autoridad / enfoque'],
[['Project Sponsor','Alex','Decisiones estratégicas, aprobación CAPEX, cliente ancla, suministro mínimo'],
 ['Project Manager','Raúl','Ejecución, cronograma, presupuesto, riesgos, asesores'],
 ['Jefe de Planta / Site Manager','Ing. industrial (SENATI)','Operación diaria; también QA y HSE en Stage 1'],
 ['Warehouse Manager','Operador de montacarga 3 t','Recepción, almacenaje, inventario, trazabilidad de lotes'],
 ['Asesor Legal / Regulatorio','Abogado local','Permisos, MOU, pacto de socios, compliance'],
 ['Asesor Ambiental (EC-RS/EIA)','Outsource','Registro EC-RS, reportes ambientales para licencias'],
])
para('Reuniones: Steering Committee mensual (Alex + Raúl + asesor legal si aplica); Project Team semanal (primeros meses); '
     'revisiones de Seguridad/Calidad quincenales desde el arranque.',size=9.5)

# ---- 7 CSF ----
h('7. Factores críticos de éxito',2)
for x in ['EC-RS iniciado de inmediato — sin registro MINAM no hay operación legal.',
          'Compromiso escrito de volumen de Alex (≥ mín. operativo) — el modelo financiero depende de ello.',
          'Cierre del gap de capital (−$858) y disponibilidad de caja para el ciclo (30–45 días).',
          'Montacarga 3 t asegurado (compra/alquiler) — sin él no se puede descargar ni apilar bolsones de 1 t.',
          'Liderazgo operativo presente (Jefe de Planta) y disciplina de procesos desde el día 1.',
          'Pacto de socios y precio de transferencia documentado (riesgo SUNAT por partes vinculadas).']:
    bullet(x)

# ---- 8 RISKS ----
h('8. Riesgos principales',2)
table(['Riesgo','Prob.','Impacto','Mitigación'],
[['EC-RS no iniciado','Alta','🔴 Crítico (bloquea operación legal)','Contratar consultor ambiental; abrir SIGERSOL; Plan de Manejo de Residuos. 60–90 d.'],
 ['Gap de capital sin cerrar','Media','Alto','Subarriendo + préstamo entre socios (notaría) + negociar depósito. No firmar nave sin cerrarlo.'],
 ['Sin compromiso de volumen (Alex)','Media','Alto','Carta/contrato de volumen mínimo (t/mes, 12 meses) antes de operar.'],
 ['Brecha de equipo (montacarga)','Media','Alto','Presupuestar alquiler ($800–1,500/mes) o unidad usada ($10–20K). Pallet jacks no manejan bolsones de 1 t.'],
 ['Retraso de permisos (EIA/ITSE)','Alta','Crítico','Asesor experto, trámites en paralelo, buffer 6–12 meses.'],
 ['Concentración en cliente único (Alex)','Media','Alto','Explorar 1–2 clientes adicionales; formalizar forecast.'],
 ['Precio de transferencia (SUNAT)','Media','Alto','Precio de mercado verificable con facturas; documentar en pacto de socios.'],
 ['Inseguridad / extorsión en zona','Media','Medio','Ubicación cuidadosa, seguridad (CCTV/vigilancia), seguro integral.'],
])

# ---- 9 AUTH ----
h('9. Autorización y aprobaciones',2)
para('Este Project Charter v2.0 autoriza la preparación y ejecución del Stage 1 (piloto de consolidación) del proyecto Factory "Amaru I" / PlasHub, con las cifras validadas indicadas.')
table(['Rol','Nombre','Firma','Fecha'],
[['Project Sponsor','Alex','_____________________','____________'],
 ['Project Manager','Raúl','_____________________','____________'],
 ['Testigo / Asesor','—','_____________________','____________'],
])

doc.save(OUT)
print('saved',OUT)
