# -*- coding: utf-8 -*-
"""Genera PlasHub_Manual_de_Uso.docx — manual de uso de la app (separado de la app).
Mismo contenido que tenía la antigua manual.html. Regenerar: python Context/_build_manual.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK   = RGBColor(0x25, 0x2A, 0x31)
ORANGE = RGBColor(0xDF, 0x69, 0x01)
TEAL   = RGBColor(0x15, 0x61, 0x6D)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREY   = RGBColor(0x6A, 0x72, 0x80)

OUT = os.path.join(os.path.dirname(__file__), "..", "PlasHub_Manual_de_Uso.docx")

doc = Document()
# márgenes
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.8)
    s.left_margin = s.right_margin = Inches(0.8)

# estilo base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)

def set_cell(cell, text, bold=False, color=None, white=False, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text); run.bold = bold; run.font.size = Pt(size)
    if white: run.font.color.rgb = WHITE
    elif color: run.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = DARK
    # borde inferior naranja
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12"); bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "DF6901")
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def h2(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = TEAL
    return p

def para(text, size=10.5, italic=False, color=None, space=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text); r.font.size = Pt(size); r.italic = italic
    if color: r.font.color.rgb = color
    return p

def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(1)
        r = p.add_run(it); r.font.size = Pt(10)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        set_cell(hdr[i], htext, bold=True, white=True, size=9)
        shade(hdr[i], "252A31")
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            set_cell(cells[i], val, size=9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def callout(text, kind="teal"):
    fill = {"teal":"E7F0F1","warn":"FDF4E3","bad":"FBE9E7","ok":"E8F5EE"}[kind]
    txtc = {"teal":RGBColor(0x10,0x46,0x4F),"warn":RGBColor(0x8A,0x5C,0x08),"bad":RGBColor(0x9A,0x2D,0x20),"ok":RGBColor(0x1E,0x6B,0x40)}[kind]
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; shade(c, fill); set_cell(c, text, color=txtc, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---------------- PORTADA / TÍTULO ----------------
tp = doc.add_paragraph(); tp.paragraph_format.space_after = Pt(2)
r = tp.add_run("Manual de uso de la app PlasHub"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = DARK
para("Cómo usar cada parámetro editable, qué afecta y cómo leer los resultados · Preparado para la reunión de socios",
     size=11, color=GREY, space=1)
para("Neo Amaru Global S.A.C. — Factory “Amaru I” · Hub de Consolidación de Plásticos · Lima Norte, Perú",
     size=10, color=GREY, space=6)
callout("Cada parámetro editable de la app tiene además un ícono (i) — pasa el cursor por encima para ver la misma "
        "explicación que aquí, sin salir de la pantalla. Punto de entrada de la app: index.html (landing) o plashub-app.html.", "teal")

# ---------------- 1 ----------------
h1("1 · Qué es la app y cómo se guarda")
para("La app es una herramienta interactiva con cuatro pestañas. Todo lo que tiene fondo crema o un control "
     "deslizante es editable: al cambiarlo, los tableros, gráficos y veredictos se recalculan al instante.")
table(["Pestaña", "Para qué sirve", "Quién la usa"],
      [["Etapa 1 · Distribución del Almacén", "Plano a escala + capacidad del hub de consolidación (cuántas toneladas aguanta la nave).", "Raúl (operación)"],
       ["Etapa 2 · Distribución de Planta", "Plano + capacidad de la planta de procesamiento (fase futura, mayor escala).", "Raúl (crecimiento)"],
       ["Plan de Negocio (E1 + E2)", "El centro de la reunión. Modelo financiero: precios, costos, capital, flujo de caja, GO/NO-GO.", "Ambos"],
       ["Tareas y Permisos", "Lista maestra de lo que falta (permisos, operación, finanzas) con avance.", "Ambos"]],
      widths=[1.9, 3.4, 1.3])
callout("IMPORTANTE — los cambios se guardan en el navegador de esta computadora (no en la nube). En otra máquina no se ven "
        "tus ediciones. Para llevar tu configuración, usa Exportar config y luego Importar.", "teal")

# ---------------- 2 ----------------
h1("2 · Controles globales")
para("Aparecen en la barra crema superior de cada herramienta:")
table(["Control", "Qué hace"],
      [["Autoguardado", "Cada edición se guarda sola en este navegador (verás “✓ guardado”). No hay botón de guardar."],
       ["Exportar config", "Descarga un .json con TODOS tus números actuales. Respaldo o para compartir."],
       ["Importar", "Carga un .json exportado y reemplaza la configuración actual."],
       ["Restablecer", "Borra tus cambios y vuelve a los valores por defecto (pide confirmación)."],
       ["Casilla “precios verificados” (Plan de Negocio)", "Mientras esté sin marcar muestra aviso BORRADOR: los precios son referencias, no cotizaciones. Para la reunión, mantenla sin marcar y dilo abiertamente."]],
      widths=[2.0, 4.6])

# ---------------- 3 ----------------
h1("3 · Guion sugerido para la reunión (~15 min)")
para("Abre la pestaña Plan de Negocio y recorre de arriba hacia abajo: la app ya está ordenada como una conversación.")
steps = [
    "Panel de Decisión GO/NO-GO (arriba del todo). Empieza aquí: “estas son las condiciones para firmar la nave”. Marca las casillas manuales según la realidad de hoy.",
    "Resumen ejecutivo. Margen, volumen de equilibrio, EBITDA del caso Base y necesidad de capital.",
    "Materiales y precios (§3). De dónde sale el margen. Aclara que los precios son estimados.",
    "Slider de volumen (§4). La palanca estrella: arrástralo entre 60 y 120 t/mes y muestra en vivo EBITDA, pozo de caja y break-even.",
    "Estrés. Baja el precio de venta a -15% y abre “Estrés de cobro” (Alex se retrasa en M3 y M4). Enseña el peor caso.",
    "Brecha de capital (§7). Edita el timing de aportes (Alex M0 / Raúl M1) y prueba el prepago de Alex. Aquí se ve que el problema real es el capital de trabajo.",
    "Etapa 2 (§10). Cierra con la visión: cambia el mes de lanzamiento y muestra el salto de EBITDA del procesamiento.",
]
for i, s in enumerate(steps, 1):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    rn = p.add_run(f"{i}. "); rn.bold = True; rn.font.color.rgb = ORANGE; rn.font.size = Pt(10.5)
    r2 = p.add_run(s); r2.font.size = Pt(10)
callout("Consejo: ten listo el .json exportado antes de la reunión, por si quieres volver a un escenario exacto o presentar desde otra computadora.", "ok")

# ---------------- 4 ----------------
h1("4 · Plan de Negocio — todos los editables")
h2("§3 · Materiales y precios")
table(["Parámetro", "Qué es / cómo se usa", "A qué afecta"],
      [["Material / Frecuencia", "Nombre del polímero y si es de alta rotación (frecuente) o no. Renombra, añade o elimina.", "Color/agrupación; en el layout E1 cambia la densidad de apilado."],
       ["Mix %", "Participación de cada material en el tonelaje. No necesita sumar 100 (se normaliza).", "Pondera los precios promedio → margen ponderado."],
       ["Compra $/t", "Precio al que compras ese material por tonelada.", "↑ compra = ↓ margen, ↓ EBITDA, pozo más profundo."],
       ["Venta $/t", "Precio al que despachas el lote consolidado.", "Sube ingreso y margen. El supuesto MÁS sensible."]],
      widths=[1.5, 3.0, 2.1])
h2("§4 · Volumen y escenarios")
table(["Parámetro", "Qué es / cómo se usa", "A qué afecta"],
      [["Slider de volumen (t/mes)", "La palanca #1. Toneladas mensuales en estado estable. Anula el volumen del escenario.", "Recalcula en vivo EBITDA, pozo de caja, break-even. Bajo ~82 t el EBITDA es negativo."],
       ["Escenarios predefinidos", "Atajos a Conservador 40 / Base 100 / Optimista 200 t.", "Al hacer clic fijan el slider y el setup CAPEX/OPEX."],
       ["Presets de rampa", "Curva de arranque: % del objetivo que logras en M1-M2-M3-M4+.", "Qué tan rápido subes a régimen. Avisa si M1 > 35%."],
       ["Setup CAPEX/OPEX (A/B)", "Lean (A) ~900 m² subarriendo · Full (B) 1,500 m² lease.", "Cambia qué tablas de CAPEX (§5) y OPEX (§6) se usan."]],
      widths=[1.5, 3.0, 2.1])
h2("§4 · Sliders de supuestos")
table(["Slider", "Qué es", "A qué afecta"],
      [["Capital confirmado", "Capital total disponible (el timing se edita en §7).", "Referencia de la brecha y del break-even de caja."],
       ["Estrés precio venta", "Sube/baja el precio de venta respecto a §3.", "Ingreso, margen, EBITDA. Peor caso."],
       ["Estrés precio compra", "Sube/baja el precio de compra.", "↑ compra = ↓ margen."],
       ["Ajuste OPEX fijo", "Escala en bloque todo el OPEX de §6.", "EBITDA y volumen de equilibrio."],
       ["Costo variable $/t", "Costo que escala con cada tonelada (flete + drayage).", "Margen de contribución; sube el equilibrio."],
       ["Desfase de cobro", "Meses entre despachar y cobrar.", "No cambia el P&L, pero profundiza el pozo de caja."],
       ["Desfase de pago a proveedor", "Meses de crédito de los proveedores.", "Pagar más tarde ALIVIA el pozo de caja."],
       ["Meses de setup (M0)", "Meses de obra/permisos antes de la 1ª tonelada.", "Más setup = pozo más profundo y temprano."],
       ["Estrés de cobro (plegable)", "Meses en que Alex se retrasa y cuántos meses extra.", "Profundiza el pozo; muestra riesgo de cliente único."]],
      widths=[1.7, 2.7, 2.2])
h2("§5-§6 · CAPEX y OPEX")
table(["Parámetro", "Qué es", "A qué afecta"],
      [["Líneas de CAPEX", "Inversión de una vez (garantía, báscula, permisos). Editable.", "CAPEX total → necesidad de capital; sale en M0."],
       ["Capital de trabajo (CT)", "Reserva para financiar el pozo M1-M3.", "Suma a la necesidad total de capital."],
       ["Líneas de OPEX", "Gasto fijo mensual (renta, sueldos, vigilancia).", "↑ OPEX = ↑ equilibrio, ↓ EBITDA."]],
      widths=[1.7, 2.9, 2.0])
h2("§7 · Brecha de capital")
table(["Parámetro", "Qué es", "A qué afecta"],
      [["Aportes de socios (nombre/monto/mes)", "Quién aporta, cuánto y CUÁNDO. Mes 0 = disponible al inicio; mes 1 = llega en M1.", "El timing cambia el pozo aunque el total sea igual: si el dinero llega tarde, el pozo se profundiza."],
       ["Prepago de Alex (%)", "% del valor de un contenedor (~20 t × precio) que Alex paga por adelantado en M0.", "Entra como caja al inicio y reduce el pozo. La palanca más rápida."]],
      widths=[2.0, 2.6, 2.0])
h2("§10 · Etapa 2 (procesamiento)")
table(["Parámetro", "Qué es", "A qué afecta"],
      [["Mes de lanzamiento / Horizonte", "Cuándo inviertes el CAPEX de E2 y migras a pellet.", "Antes = más margen pero más capital antes."],
       ["E1 · Margen neto ($/t)", "Se sincroniza desde el modelo E1 (botón ↺ Sync E1).", "EBITDA de E1 mientras esperas a lanzar E2."],
       ["E2 · Volumen / Rampa / Merma", "Toneladas objetivo, meses para llegar y % de pérdida.", "Ingreso y EBITDA de la fase E2."],
       ["E2 · Precio pellet / Costo MP", "Precio del pellet ($600-1,300/t) y costo del material.", "El gran salto de valor de la E2."],
       ["Costo de procesamiento variable", "Energía, agua, mano de obra, empaque — por tonelada.", "Se suma al costo de materia prima."],
       ["CAPEX / OPEX de E2 (tablas)", "Equipos/obra (una vez) y gasto mensual de la planta.", "Capital de lanzamiento y OPEX desde el arranque."]],
      widths=[2.0, 2.6, 2.0])

# ---------------- 5 ----------------
h1("5 · Distribución Etapa 1 — todos los editables")
callout("Idea clave: el almacén se dimensiona por FUNCIÓN. Las zonas de soporte (oficina, pesaje) se calculan por su driver "
        "(personas, equipo) y NO crecen con la nave; el espacio que sobra se vuelve “buffer libre”. "
        "Capacidad = mín(límite de almacenamiento, límite de manejo).", "teal")
table(["Grupo", "Parámetros", "A qué afecta"],
      [["Huella de la nave", "Frente, profundidad, altura libre.", "Área total → más área = más buffer libre. Avisa si la nave es muy corta."],
       ["Materiales y equipo", "Taxonomía, mix por material; montacarga vs transpaleta.", "El mix dimensiona el almacén por polímero. Con transpaleta la capacidad colapsa."],
       ["Supuestos de operación", "Montacargas, horas-turno, lead time, movimientos/t, ciclo, utilización.", "Capacidad de MANEJO (t/mes) — el cuello de botella real."],
       ["Densidades y días de hold", "Toneladas por m² y días de inventario objetivo por material.", "Buffer = mix% × días × volumen ÷ densidad. ↑ días = más área."],
       ["Zonas funcionales", "Staff, báscula, spots de clasificación, recepción.", "Tamaño de oficina/pesaje/clasificación — independiente de la nave."],
       ["m² editables en la tabla", "Sobrescribe a mano un área; ↺ vuelve al valor calculado.", "Solo esa zona; el resto se recalcula."]],
      widths=[1.7, 2.9, 2.0])

# ---------------- 6 ----------------
h1("6 · Distribución Etapa 2 — todos los editables")
callout("Aquí el cuello de botella es la LÍNEA de procesamiento (no el montacarga). "
        "Capacidad = líneas × t/h × horas × turnos × días × utilización. Las zonas de soporte (oficina, lab QA, PTAR) "
        "se dimensionan por su driver y no escalan con la nave.", "teal")
table(["Grupo", "Parámetros", "A qué afecta"],
      [["Huella de la nave", "Frente, fondo, altura, objetivo t/mes.", "Área total y dimensionado de buffers/despacho."],
       ["Línea de procesamiento", "Líneas, t/h, horas/turno, turnos, días, utilización, merma.", "Definen la CAPACIDAD mensual y las toneladas vendibles."],
       ["Almacenamiento (buffers)", "Densidad y días de buffer de MP y PT.", "Área de los almacenes (no el throughput)."],
       ["Zonas funcionales", "Supervisores, técnicos, familias, agua m³/t, PTAR, áreas mínimas.", "Oficina, laboratorio, planta de tratamiento de agua, recepción/despacho."],
       ["m² editables en la tabla", "Sobrescribe a mano un área de soporte; ↺ al valor funcional.", "Solo esa zona."]],
      widths=[1.7, 2.9, 2.0])

# ---------------- 7 ----------------
h1("7 · Tareas y Permisos")
table(["Elemento", "Qué hace"],
      [["Casilla de cada tarea", "Marca lo completado. Se guarda y se SINCRONIZA con la sección “Permisos” del Plan de Negocio."],
       ["Filtros (Etapa/Tipo/Prioridad/Estado)", "Ocultan lo que no coincide (no borran). Combínalos."],
       ["Reiniciar", "Vuelve todas las casillas a su estado inicial."]],
      widths=[2.2, 4.4])
callout("Ruta crítica: EC-RS (MINAM/SIGERSOL) y EIA son los trámites largos (60-180 días). Empieza por ahí — sin EC-RS no se puede operar legalmente.", "warn")

# ---------------- 8 ----------------
h1("8 · Cómo leer los resultados")
table(["Salida", "Qué te dice"],
      [["Panel GO/NO-GO", "Condiciones para firmar la nave. (i) = calculadas; (mano) = las marcas tú. Verde=listo, amarillo=condicional, rojo=no avances."],
       ["EBITDA (estable)", "Ganancia operativa mensual en régimen, tras OPEX. Negativo = bajo el equilibrio."],
       ["Equilibrio (P&L)", "Volumen mínimo para EBITDA = 0. Por encima ganas ~$112 por tonelada extra."],
       ["Break-even de CAJA", "Volumen mínimo para que el banco NUNCA quede negativo. Si dice “no alcanzable”, con el capital actual ningún volumen evita el pozo: el problema es de capital. Sube el “Capital confirmado” y aparecerá un número."],
       ["Pozo de caja", "El punto más bajo del banco y en qué mes = cuánto dinero necesitas para no quebrar antes de cobrar. La cifra clave para financiamiento."],
       ["Matriz de sensibilidad", "EBITDA según precio × volumen. Verde=positivo, amarillo=al límite, rojo=pérdida."],
       ["Gráfico de caja (M0→M12)", "La curva del banco. M0 = setup. La línea roja es el $0: si baja de ahí, insolvencia sin financiamiento."],
       ["Tabla P&L (Ejecutiva/Detallada)", "Ejecutiva = 4 columnas clave (ideal para presentar). Detallada = los 12 meses."],
       ["Tablero de capacidad (E1/E2)", "Cuántas t/mes aguanta la instalación y si la limita el almacenamiento o el manejo/línea."]],
      widths=[2.0, 4.6])

# ---------------- 9 ----------------
h1("9 · Glosario")
gloss = [
    ("CAPEX", "Inversión de una vez (equipos, garantías, permisos). Sale al inicio."),
    ("OPEX", "Gasto fijo mensual (renta, sueldos, servicios) que pagas opere o no."),
    ("EBITDA", "Ganancia operativa antes de intereses, impuestos y depreciación. Aquí ≈ ingreso − costos variables − OPEX fijo."),
    ("Margen de contribución ($/t)", "Lo que aporta cada tonelada tras restar compra y costo variable. Cubre el OPEX fijo; lo que sobra es EBITDA."),
    ("Capital de trabajo (CT)", "Dinero para financiar el desfase entre pagar (ahora) y cobrar (después). La verdadera brecha de PlasHub."),
    ("Pozo de caja", "El saldo más bajo que alcanza el banco. Define cuánto capital necesitas tener listo."),
    ("Equilibrio P&L vs CAJA", "P&L: volumen donde el EBITDA es 0. CAJA: volumen donde el banco nunca es negativo (más alto y más relevante)."),
    ("Desfase de cobro / pago (lag)", "Meses entre la venta y el cobro, o entre la compra y el pago al proveedor."),
    ("Rampa", "Qué % del volumen objetivo logras cada mes al arrancar (no se llega al 100% el primer mes)."),
    ("Merma", "% de material que se pierde en el proceso (Etapa 2). Reduce las toneladas vendibles."),
    ("Throughput vs buffer", "Throughput = toneladas por mes (límite de manejo/línea). Buffer = inventario en piso (días)."),
    ("Driver funcional", "La variable que dimensiona una zona (personas para la oficina, caudal para la PTAR) en vez de un % del área."),
    ("EC-RS", "Registro de Empresa Comercializadora de Residuos Sólidos (MINAM/SIGERSOL). Permiso de ruta crítica para operar legalmente."),
]
for term, desc in gloss:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(term + ": "); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK
    r2 = p.add_run(desc); r2.font.size = Pt(10)

# ---------------- 10 ----------------
h1("10 · Lo que NO está confirmado (dilo en la reunión)")
callout("Precios sin verificar. Compra y venta son referencias trianguladas, no cotizaciones firmes. La casilla "
        "“precios verificados” está apagada a propósito. Acción: conseguir 2-3 cotizaciones reales antes de levantar capital.", "warn")
callout("La brecha real es capital de trabajo. Aunque el P&L sea positivo, el ciclo comprar-ahora/cobrar-después lleva el "
        "banco a negativo en los primeros meses. Por eso el break-even de caja puede salir “no alcanzable” con $22K: hay "
        "que cerrar el pozo con prepago de Alex, garantía a 1 mes y/o puente entre socios.", "bad")
callout("Brecha de equipo (Etapa 1). El montacarga contrabalanceado de 3 t es indispensable; con transpaleta la capacidad "
        "colapsa. Hay que presupuestar su alquiler o compra.", "bad")
callout("Etapa 2 = supuestos gruesos. Precio de pellet, CAPEX de equipos, merma y OPEX son de orden de magnitud. Sirve para "
        "la visión y el “cuándo lanzar”, no como cifras firmes.", "warn")
callout("SUNAT / precio de transferencia. Alex es proveedor + cliente + socio: las facturas deben ir a precio de mercado y "
        "documentarse, y conviene un pacto de socios notariado.", "teal")

doc.save(os.path.abspath(OUT))
print("OK ->", os.path.abspath(OUT))
